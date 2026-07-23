"""
成果导出服务（阶段 8E）。

只导出白名单字段（成果标题、来源任务安全摘要、店铺名称、状态、
生成时间、结构化正文），不导出：

- system prompt / Provider 原始响应 / API Key；
- 完整内部 Context；
- 店铺凭据（明文或密文）；
- Task.payload 原文。

全部在内存中生成（io.BytesIO），不写入磁盘、不产生需要清理的
临时文件，也就不存在目录穿越风险——导出参数只有 deliverable_id
（整数，由路由层的 Pydantic 路径参数校验）和 format（受限枚举），
不接受任意文件路径。
"""

import io
import json
from datetime import datetime
from urllib.parse import quote

from app.models.deliverable_db import DeliverableDB, DeliverableVersionDB
from app.services.deliverable_content import (
    DELIVERABLE_TYPE_TITLE_PREFIX,
    render_markdown_body,
)

_STATUS_LABELS = {
    "draft": "草稿",
    "pending_review": "待审核",
    "approved": "已批准",
    "rejected": "已驳回",
    "converted_to_task": "已转任务",
    "archived": "已归档",
}

SUPPORTED_FORMATS = ("markdown", "json", "pdf", "docx", "xlsx")


class UnsupportedExportFormatError(Exception):
    pass


def _status_label(status: str) -> str:
    return _STATUS_LABELS.get(status, status)


def _deliverable_type_label(deliverable_type: str) -> str:
    return DELIVERABLE_TYPE_TITLE_PREFIX.get(deliverable_type, deliverable_type)


def build_export_filename(deliverable: DeliverableDB, extension: str) -> tuple[str, str]:
    """
    返回 (ascii_fallback_filename, utf8_filename)，供路由层拼接
    安全的 Content-Disposition 头（同时提供 filename 与
    filename*，兼容不支持 RFC 5987 的客户端）。
    """

    ascii_name = f"{deliverable.deliverable_code}.{extension}"
    utf8_name = f"{deliverable.title}-{deliverable.deliverable_code}.{extension}"
    return ascii_name, utf8_name


def build_content_disposition(ascii_name: str, utf8_name: str) -> str:
    quoted = quote(utf8_name)
    return f'attachment; filename="{ascii_name}"; filename*=UTF-8\'\'{quoted}'


def _meta_lines(deliverable: DeliverableDB, shop_name: str | None) -> list[str]:
    return [
        f"- 来源任务：{deliverable.source_task_id}",
        f"- AI 员工：{deliverable.agent_name}",
        f"- 店铺：{shop_name or '未绑定店铺'}",
        f"- 当前状态：{_status_label(deliverable.status)}",
        f"- 当前版本：v{deliverable.current_version}",
        f"- 生成时间：{deliverable.created_at.isoformat() if deliverable.created_at else ''}",
    ]


def export_markdown(
    deliverable: DeliverableDB, version: DeliverableVersionDB, shop_name: str | None
) -> bytes:
    lines = [f"# {deliverable.title}", ""]
    lines += _meta_lines(deliverable, shop_name)
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append(version.content)
    return "\n".join(lines).encode("utf-8")


def export_json(
    deliverable: DeliverableDB, version: DeliverableVersionDB, shop_name: str | None
) -> bytes:
    payload = {
        "deliverable_code": deliverable.deliverable_code,
        "title": deliverable.title,
        "deliverable_type": deliverable.deliverable_type,
        "status": deliverable.status,
        "source_task_id": deliverable.source_task_id,
        "root_task_id": deliverable.root_task_id,
        "shop_name": shop_name,
        "agent_name": deliverable.agent_name,
        "current_version": deliverable.current_version,
        "created_at": deliverable.created_at.isoformat() if deliverable.created_at else None,
        "version": {
            "version_number": version.version_number,
            "format": version.format,
            "structured_content": version.structured_content,
        },
    }
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def export_pdf(
    deliverable: DeliverableDB, version: DeliverableVersionDB, shop_name: str | None
) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    font_name = "STSong-Light"
    if font_name not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DeliverableTitle", parent=styles["Title"], fontName=font_name, fontSize=18
    )
    meta_style = ParagraphStyle(
        "DeliverableMeta", parent=styles["Normal"], fontName=font_name, fontSize=10,
        textColor="#555555",
    )
    body_style = ParagraphStyle(
        "DeliverableBody", parent=styles["Normal"], fontName=font_name, fontSize=11,
        leading=16,
    )
    heading_style = ParagraphStyle(
        "DeliverableHeading", parent=styles["Heading3"], fontName=font_name, fontSize=13,
    )

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
    )

    story = [Paragraph(deliverable.title, title_style), Spacer(1, 8)]

    for line in _meta_lines(deliverable, shop_name):
        story.append(Paragraph(line.replace("- ", ""), meta_style))
    story.append(Spacer(1, 12))

    for raw_line in version.content.split("\n"):
        line = raw_line.strip()
        if not line:
            story.append(Spacer(1, 4))
            continue
        if line.startswith("### "):
            story.append(Paragraph(line[4:], heading_style))
        elif line.startswith("- "):
            story.append(Paragraph(f"• {line[2:]}", body_style))
        else:
            story.append(Paragraph(line, body_style))

    doc.build(story)
    return buffer.getvalue()


def export_docx(
    deliverable: DeliverableDB, version: DeliverableVersionDB, shop_name: str | None
) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading(deliverable.title, level=1)

    for line in _meta_lines(deliverable, shop_name):
        document.add_paragraph(line.replace("- ", ""))

    document.add_paragraph("")

    for raw_line in version.content.split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(line)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _autosize_columns(worksheet) -> None:
    for column_cells in worksheet.columns:
        length = max((len(str(cell.value)) if cell.value else 0) for cell in column_cells)
        worksheet.column_dimensions[column_cells[0].column_letter].width = min(
            max(length + 2, 12), 60
        )


def _write_list_sheet(workbook, title: str, header: str, items) -> None:
    sheet = workbook.create_sheet(title=title[:31])
    sheet.append([header])
    for item in items or []:
        if isinstance(item, dict):
            sheet.append([json.dumps(item, ensure_ascii=False)])
        else:
            sheet.append([str(item)])
    _autosize_columns(sheet)


def export_xlsx(
    deliverable: DeliverableDB, version: DeliverableVersionDB, shop_name: str | None
) -> bytes:
    from openpyxl import Workbook

    structured = version.structured_content or {}

    workbook = Workbook()
    summary_sheet = workbook.active
    summary_sheet.title = "摘要"
    summary_sheet.append(["字段", "内容"])
    summary_sheet.append(["标题", deliverable.title])
    summary_sheet.append(["AI 员工", deliverable.agent_name])
    summary_sheet.append(["店铺", shop_name or "未绑定店铺"])
    summary_sheet.append(["状态", _status_label(deliverable.status)])
    summary_sheet.append(["当前版本", f"v{deliverable.current_version}"])
    summary_sheet.append([
        "生成时间",
        deliverable.created_at.isoformat() if deliverable.created_at else "",
    ])
    summary_sheet.append(["摘要", structured.get("summary", "") if version.format == "structured" else structured.get("text", "")])
    _autosize_columns(summary_sheet)

    if version.format != "structured":
        buffer = io.BytesIO()
        workbook.save(buffer)
        return buffer.getvalue()

    if deliverable.deliverable_type == "ceo_analysis":
        _write_list_sheet(workbook, "发现", "findings", structured.get("findings"))
        _write_list_sheet(workbook, "风险", "risks", structured.get("risks"))
        _write_list_sheet(workbook, "行动", "actions", structured.get("actions"))
        _write_list_sheet(workbook, "委派", "delegations", structured.get("delegations"))

    elif deliverable.deliverable_type == "sales_analysis":
        _write_list_sheet(workbook, "已知事实", "known_facts", structured.get("known_facts"))
        _write_list_sheet(workbook, "数据缺口", "data_gaps", structured.get("data_gaps"))
        _write_list_sheet(workbook, "机会", "opportunities", structured.get("opportunities"))
        _write_list_sheet(workbook, "策略", "strategy", [structured.get("strategy")] if structured.get("strategy") else [])
        _write_list_sheet(workbook, "输入要求", "required_inputs", structured.get("required_inputs"))
        _write_list_sheet(workbook, "风险", "warnings", structured.get("warnings"))

    elif deliverable.deliverable_type == "product_analysis":
        _write_list_sheet(workbook, "推荐结论", "selection_verdict", [structured.get("selection_verdict")] if structured.get("selection_verdict") else [])
        _write_list_sheet(workbook, "商品机会", "opportunities", structured.get("opportunities"))
        _write_list_sheet(workbook, "商品组合", "assortment_plan", [structured.get("assortment_plan")] if structured.get("assortment_plan") else [])
        _write_list_sheet(workbook, "最小测试", "minimum_viable_test", [structured.get("minimum_viable_test")] if structured.get("minimum_viable_test") else [])
        _write_list_sheet(workbook, "上架清单", "listing_checklist", structured.get("listing_checklist"))
        _write_list_sheet(workbook, "供应商问题", "supplier_questions", structured.get("supplier_questions"))
        _write_list_sheet(workbook, "下一步", "next_actions", structured.get("next_actions"))
        _write_list_sheet(workbook, "风险", "warnings", structured.get("warnings"))

    else:
        for key, value in structured.items():
            if key == "summary":
                continue
            _write_list_sheet(
                workbook, str(key), str(key), value if isinstance(value, list) else [value]
            )

    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


_EXPORTERS = {
    "markdown": (export_markdown, "md", "text/markdown; charset=utf-8"),
    "json": (export_json, "json", "application/json"),
    "pdf": (export_pdf, "pdf", "application/pdf"),
    "docx": (
        export_docx,
        "docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ),
    "xlsx": (
        export_xlsx,
        "xlsx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ),
}


def export_deliverable(
    deliverable: DeliverableDB,
    version: DeliverableVersionDB,
    export_format: str,
    shop_name: str | None,
) -> tuple[bytes, str, str, str]:
    """
    返回 (文件字节内容, content_type, ascii_filename, utf8_filename)。
    """

    if export_format not in _EXPORTERS:
        raise UnsupportedExportFormatError(export_format)

    exporter, extension, content_type = _EXPORTERS[export_format]
    content = exporter(deliverable, version, shop_name)
    ascii_name, utf8_name = build_export_filename(deliverable, extension)

    return content, content_type, ascii_name, utf8_name
